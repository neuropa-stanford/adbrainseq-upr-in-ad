#!/usr/bin/env python3
"""Supplementary Table 4 — SEA-AD (Allen Institute) clinicopathologic information for the 84 donors
used in the independent-replication analysis (Figure 6). Style is an exact clone of S.Table 2 in
'Supplementary table 1 2 3_AD brain tissue information.docx': Arial 11 pt, every cell centered,
single 0.5 pt (sz=4) cell borders on all four sides, header row shaded E7E6E6, fixed layout.
Non-destructive: writes a NEW standalone .docx; the original 1-2-3 file is untouched."""
import csv, openpyxl, copy
import docx
from docx.shared import Pt, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "/data/adbrainseq/anc"
RAW = BASE + "/Major Revision/processed raw data"

# ---- gather data (analysis donors == pseudobulk donors) ----
used = sorted(set(r["donor"] for r in csv.DictReader(open(RAW + "/SEAAD_3group_donor_pseudobulk.csv"))))
wb = openpyxl.load_workbook(RAW + "/SEAAD_donor_metadata_SuppTable1.xlsx", read_only=True); ws = wb.active
h = [str(c) for c in next(ws.iter_rows(min_row=1, max_row=1, values_only=True))]
I = lambda k: next(i for i, c in enumerate(h) if k.lower() in c.lower())
c = {k: I(k) for k in ["Donor ID", "Age at Death", "Sex", "APOE", "Cognitive Status",
                       "Thal", "Braak", "CERAD", "Overall AD neuropath", "PMI"]}
meta = {r[c["Donor ID"]]: r for r in ws.iter_rows(min_row=2, values_only=True) if r[c["Donor ID"]]}
wb.close()

BRK = {"Braak 0": 0, "Braak I": 1, "Braak II": 2, "Braak III": 3, "Braak IV": 4, "Braak V": 5, "Braak VI": 6}
def pmi(v):
    try: return "%.1f" % float(v)
    except (TypeError, ValueError): return "n/a"
recs = []
for d in used:
    r = meta[d]
    recs.append([str(r[c["Donor ID"]]), str(r[c["Age at Death"]]), str(r[c["Sex"]]), str(r[c["APOE"]]),
                 str(r[c["Cognitive Status"]]), str(r[c["Thal"]]), str(r[c["Braak"]]), str(r[c["CERAD"]]),
                 str(r[c["Overall AD neuropath"]]), pmi(r[c["PMI"]])])
recs.sort(key=lambda x: (BRK.get(x[6], 99), x[0]))   # by Braak stage, then Donor ID

HEADERS = ["Donor ID", "Age", "Sex", "APOE", "Cognitive status", "Thal phase",
           "Braak stage", "CERAD score", "ADNC", "PMI (h)"]
# column widths in twips (fixed layout), tuned to content; sums ~= landscape text width
WIDTHS = [1500, 620, 720, 760, 1500, 1080, 1180, 1200, 1180, 820]

# ---- build document ----
doc = docx.Document()
doc.styles["Normal"].font.name = "Arial"
doc.styles["Normal"].font.size = Pt(11)
sec = doc.sections[0]
from docx.enum.section import WD_ORIENT
sec.orientation = WD_ORIENT.LANDSCAPE
sec.page_width, sec.page_height = sec.page_height, sec.page_width
sec.left_margin = sec.right_margin = Twips(720)

# caption (bold "S. Table 4." lead, matching S.Table 1-3 wording)
cap = doc.add_paragraph()
r1 = cap.add_run("S. Table 4. ")
r1.bold = True
cap.add_run("Clinicopathologic information for the SEA-AD (Allen Institute for Brain Science) "
            "cohort used in the independent replication analysis (Figure 6). All 84 donors "
            "(Middle temporal gyrus, snRNA-seq) are listed with age at death, sex, APOE genotype, "
            "cognitive status, and AD neuropathologic staging (Thal phase, Braak stage, CERAD "
            "neuritic-plaque score, and the integrated ADNC level), sorted by Braak stage. "
            "PMI, post-mortem interval in hours.")
for run in cap.runs:
    run.font.name = "Arial"; run.font.size = Pt(11)

def set_borders(tc):
    tcPr = tc.get_or_add_tcPr()
    b = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        e = OxmlElement("w:" + side)
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), "auto")
        b.append(e)
    tcPr.append(b)

def shade(tc, fill):
    tcPr = tc.get_or_add_tcPr()
    s = OxmlElement("w:shd")
    s.set(qn("w:val"), "clear"); s.set(qn("w:color"), "auto"); s.set(qn("w:fill"), fill)
    tcPr.append(s)

def vcenter(tc):
    tcPr = tc.get_or_add_tcPr()
    v = OxmlElement("w:vAlign"); v.set(qn("w:val"), "center"); tcPr.append(v)

def fill_cell(cell, text, bold=False, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text); run.font.name = "Arial"; run.font.size = Pt(11); run.bold = bold
    vcenter(cell._tc); set_borders(cell._tc)
    if fill: shade(cell._tc, fill)

table = doc.add_table(rows=1 + len(recs), cols=len(HEADERS))
table.alignment = WD_TABLE_ALIGNMENT.CENTER
# fixed layout + grid widths
tblPr = table._tbl.tblPr
lay = OxmlElement("w:tblLayout"); lay.set(qn("w:type"), "fixed"); tblPr.append(lay)
grid = table._tbl.find(qn("w:tblGrid"))
for gc, w in zip(grid.findall(qn("w:gridCol")), WIDTHS):
    gc.set(qn("w:w"), str(w))

for j, htxt in enumerate(HEADERS):
    cell = table.rows[0].cells[j]
    cell.width = Twips(WIDTHS[j])
    fill_cell(cell, htxt, bold=False, fill="E7E6E6")
for i, rec in enumerate(recs, start=1):
    for j, val in enumerate(rec):
        cell = table.rows[i].cells[j]
        cell.width = Twips(WIDTHS[j])
        fill_cell(cell, val)

out = BASE + "/Supplementary Table 4_SEA-AD clinical information.docx"
doc.save(out)
print("wrote", out, "| rows:", len(recs), "| cols:", len(HEADERS))
# quick sanity
from collections import Counter
print("Braak:", dict(Counter(x[6] for x in recs)))
