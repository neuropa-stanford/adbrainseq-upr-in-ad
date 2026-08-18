#!/usr/bin/env python3
"""Non-destructive manuscript edits for the SEA-AD replication + IHC renumber.
Saves a NEW versioned .docx; the original 06042026 file is untouched.
  (1) Insert SEA-AD replication results subsection (heading+body) before the IHC subheading.
  (2) Insert new Fig. 6 legend (SEA-AD) before the old Fig. 6 (IHC) legend.
  (3) Renumber old Fig. 6 -> Fig. 7: legend heading, rewrite legend body for the new a..g scheme,
      and remap the in-text callouts in the IHC results paragraph.
  (4) Add Supplementary Table 4 (SEA-AD clinical) line after Supplementary Table 3.
  (5) Add a Supplementary Figures section + Supplementary Figure 1 legend (per-donor IHC montage)."""
import docx
from docx.shared import Pt

SRC = ("/data/adbrainseq/Publication/2024_scRNA seq/V1 Science manuscript/"
       "ADBrainSeq_V6.0/Acta Neuropathologica Communication/UPRinADBrainSeq_ANC_GPark_06042026.docx")
OUT = ("/data/adbrainseq/Publication/2024_scRNA seq/V1 Science manuscript/"
       "ADBrainSeq_V6.0/Acta Neuropathologica Communication/UPRinADBrainSeq_ANC_GPark_REVISION_20260814.docx")
FONT = "Times New Roman"
d = docx.Document(SRC)
P = d.paragraphs

def find(pred):
    return next(p for p in P if pred(p.text))

def add_before(ref, specs, style="Normal1", size=None):
    """specs = list of (text, bold). Inserts a new paragraph immediately before ref."""
    np = ref.insert_paragraph_before(style=style)
    for text, bold in specs:
        r = np.add_run(text); r.bold = bold; r.font.name = FONT
        if size: r.font.size = size
    return np

# ---------- anchors ----------
ihc_head = find(lambda t: t.strip().startswith("Immunohistochemical evaluation of TRIB3"))
old_fig6_head = find(lambda t: t.strip().startswith("Fig. 6 UPR pathway disruption"))
supp_tab3 = find(lambda t: t.strip().startswith("Supplementary Table 3"))
declar = find(lambda t: t.strip() == "Declarations")

# ---------- (3) remap in-text IHC callouts (all in the IHC results paragraph) ----------
ihc_body = find(lambda t: t.strip().startswith("Next, we performed immunohistochemical staining"))
for run in ihc_body.runs:
    if "Fig. 6" in run.text:
        run.text = (run.text
                    .replace("(Fig. 6a)", "(Fig. 7a)")
                    .replace("(Fig. 6b, 6e)", "(Fig. 7b, 7c)")
                    .replace("(Fig. 6c, 6f, 6d, 6g)", "(Fig. 7d, 7e, 7f, 7g)")
                    .replace("(Fig. 4d-f and 6e-g)", "(Fig. 4d-f and 7c, 7e, 7g)"))

# ---------- (3) renumber old Fig 6 heading -> Fig 7 ----------
old_fig6_head.runs[0].text = old_fig6_head.runs[0].text.replace("Fig. 6", "Fig. 7", 1)
# rewrite old Fig 6 legend BODY (paragraph right after the heading) for the new a..g scheme
fig7_body = old_fig6_head._p.getnext()  # xml element of next paragraph
fig7_body_par = next(p for p in P if p._p is fig7_body)
FIG7_LEGEND = ("(a) Phospho-Tau (AT8) immunostaining in control (Braak I) and tauopathy (Braak VI) brain "
    "tissue, confirming tau pathology. (b, d, f) Immunofluorescent double staining in control and AD brains for "
    "(b) neurons (MAP2, red; TRIB3, green; DAPI, blue), (d) microglia (CD45, red; TMED2, green; DAPI, blue), and "
    "(f) oligodendrocytes (MOG, red; TMED2, green; DAPI, blue); merged and 3.5x-magnified insets are shown. "
    "(c, e, g) Quantification of double-positive cells for (c) TRIB3 in MAP2+ neurons, (e) TMED2 in CD45+ "
    "microglia, and (g) TMED2 in MOG+ oligodendrocytes, in control versus AD brains. Brain tissues from three "
    "donors per condition were analyzed, with triplicate stainings and images per sample; per-donor images for "
    "all cases are shown in Supplementary Figure 1. Statistical significance: **p ≤ 0.01; ****p ≤ 0.0001.")
fig7_body_par.runs[0].text = FIG7_LEGEND
for extra in fig7_body_par.runs[1:]:
    extra.text = ""

# ---------- (2) insert new Fig. 6 (SEA-AD) legend before the (now) Fig. 7 legend ----------
FIG6_BODY = ("(a) Independent replication in the SEA-AD (Allen Institute for Brain Science) middle temporal "
    "gyrus snRNA-seq cohort (84 donors; Supplementary Table 4), a different platform and brain region, "
    "stratified by Braak stage (low-Braak 0-II; early-AD III-IV; late-AD V-VI). Per-gene log2 fold-change of the "
    "Response-to-ER-stress gene set (254 genes) for early-AD and late-AD versus low-Braak, per cell type "
    "(two-tone violins; red, set mean; blue, quartiles); significance stars, gene-level Wilcoxon signed-rank; "
    "the donor-level ordinal-Braak trend (Spearman, n = 84) is annotated per cell type. (b) The three canonical "
    "UPR sensor transcripts (EIF2AK3/PERK, ERN1/IRE1, ATF6) across Braak stage, donor-level; the sensors are "
    "coordinately down-regulated in neurons and show the opposite microglial trend, recapitulating the "
    "neuron/glia divergence in an independent cohort. Statistical significance: *p < 0.05; **p < 0.01; "
    "***p < 0.001; ****p < 0.0001.")
add_before(old_fig6_head, [("Fig. 6 ", True),
    ("Neuronal attenuation of the UPR/ER-stress transcriptional program reproduces in an independent "
     "SEA-AD cohort.", True)])
add_before(old_fig6_head, [(FIG6_BODY, False)])

# ---------- (1) insert SEA-AD replication results subsection before the IHC subheading ----------
SEAAD_HEAD = ("Neuronal down-regulation of UPR-associated transcripts reproduces in an independent "
              "SEA-AD cohort")
SEAAD_BODY = ("\tTo test whether the neuronal decline in UPR-associated transcription generalizes beyond our "
    "discovery cohort, we analyzed the independent SEA-AD (Allen Institute) middle temporal gyrus snRNA-seq "
    "cohort (84 Braak-staged donors; Supplementary Table 4), profiled on a different platform and brain region. "
    "Grouping donors by Braak stage (0-II, III-IV, V-VI), the Response-to-ER-stress transcriptional program "
    "shifted downward in both excitatory and inhibitory neurons with advancing tau pathology, reproducing the "
    "direction seen in the discovery cohort (Fig. 6a). Notably, the three canonical UPR sensors "
    "(EIF2AK3/PERK, ERN1/IRE1, ATF6) were themselves coordinately down-regulated in neurons — most "
    "significantly in excitatory neurons (ERN1/IRE1 ρ = −0.35, p = 0.00097; ATF6 ρ = −0.29, "
    "p = 0.0081) and inhibitory neurons (ATF6 ρ = −0.27, p = 0.014) — whereas microglia showed the "
    "opposite trend (ERN1/IRE1 ρ = +0.36, p = 0.00078), recapitulating the neuron/glia divergence in a fully "
    "independent dataset (Fig. 6b). These orthogonal, independent-cohort findings reinforce that the neuronal "
    "UPR transcriptional program is attenuated with advancing tau pathology.")
add_before(ihc_head, [(SEAAD_HEAD, True)])
add_before(ihc_head, [(SEAAD_BODY, False)])

# ---------- (4) Supplementary Table 4 after Supplementary Table 3 ----------
tab_next = supp_tab3._p.getnext()
tab_next_par = next(p for p in P if p._p is tab_next)
S4 = ("Supplementary Table 4: Clinicopathologic information for the SEA-AD (Allen Institute for Brain Science) "
      "cohort used in the independent replication analysis (Figure 6), including age, sex, APOE genotype, "
      "cognitive status, and AD neuropathologic staging (Thal phase, Braak stage, CERAD score, and ADNC), "
      "for all 84 donors.")
add_before(tab_next_par, [(S4, False)], style="Normal", size=Pt(11))

# ---------- (5) Supplementary Figures section + legend, before Declarations ----------
SF1 = ("Supplementary Figure 1: Per-donor immunohistochemical validation of TRIB3 and TMED2 across Braak stage. "
       "Immunofluorescent staining for (a) neurons (MAP2/TRIB3/DAPI), (b) microglia (CD45/TMED2/DAPI), and "
       "(c) oligodendrocytes (MOG/TMED2/DAPI) in Braak 0 versus Braak VI brain tissue, shown for all three "
       "donors per condition in triplicate (#1-#3), with merged and 3.5x insets. These per-case images "
       "correspond to the representative panels and quantification in Figure 7.")
add_before(declar, [("Supplementary Figures", True)], style="Normal", size=Pt(11))
add_before(declar, [(SF1, False)], style="Normal", size=Pt(11))

d.save(OUT)
print("saved", OUT)
